import os
from zipfile import ZipFile
from docx import Document

p = r"E:\VibeFit\docs\VibeFit_Product_Market_Research.docx"
d = Document(p)
with ZipFile(p) as z:
    broken = z.testzip()
    rels = z.read("word/_rels/document.xml.rels")
    links = rels.count(b'TargetMode="External"')

assert broken is None
assert len(d.paragraphs) > 100
assert len(d.tables) == 5
assert links >= 11
assert all(s.page_width and s.page_height for s in d.sections)
assert all(len(t.columns) in (1, 3) for t in d.tables)

print({
    "paragraphs": len(d.paragraphs),
    "tables": len(d.tables),
    "sections": len(d.sections),
    "headings": sum(1 for x in d.paragraphs if x.style.name.startswith("Heading")),
    "external_links": links,
    "bytes": os.path.getsize(p),
    "zip_test": broken,
})
