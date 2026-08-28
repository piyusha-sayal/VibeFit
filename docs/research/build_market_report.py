from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = r"E:\VibeFit\docs\VibeFit_Product_Market_Research.docx"

NAVY = RGBColor(31, 55, 72)
TEAL = RGBColor(28, 128, 126)
MUTED = RGBColor(90, 100, 110)
LIGHT = "EAF4F3"
PALE = "F4F6F8"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1C807E")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    run.append(rpr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    link.append(run)
    paragraph._p.append(link)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label.upper() + "\n")
    r.bold = True
    r.font.color.rgb = TEAL
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

for name, size, color, before, after in [
    ("Title", 28, NAVY, 0, 8),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 17, NAVY, 16, 8),
    ("Heading 2", 13, TEAL, 12, 5),
    ("Heading 3", 11, NAVY, 9, 3),
]:
    s = styles[name]
    s.font.name = "Arial"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = name != "Subtitle"
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Bullet 2", "List Number"]:
    s = styles[name]
    s.font.name = "Arial"
    s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(4)

# Running furniture
hp = section.header.paragraphs[0]
hp.text = "VIBEFIT  |  PRODUCT-MARKET RESEARCH"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in hp.runs:
    run.font.name = "Arial"; run.font.size = Pt(8); run.font.color.rgb = MUTED; run.bold = True
fp = section.footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp.add_run("VibeFit • August 2026  |  ")
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)
for run in fp.runs:
    run.font.name = "Arial"; run.font.size = Pt(8); run.font.color.rgb = MUTED

# Editorial cover
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("MARKET & PRODUCT STRATEGY")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = TEAL
p = doc.add_paragraph("VibeFit", style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph("Korean-Inspired Total Appearance Intelligence", style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("From selfie analysis to trusted routines, shopping decisions, and outcomes")
r.italic = True; r.font.name = "Arial"; r.font.size = Pt(11); r.font.color.rgb = MUTED
for _ in range(5):
    doc.add_paragraph()
p = doc.add_paragraph("Prepared for the VibeFit founder/product team\nAugust 25, 2026")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.name = "Arial"; run.font.size = Pt(9.5); run.font.color.rgb = MUTED
doc.add_page_break()

doc.add_heading("Executive answer", level=1)
add_callout(doc, "Recommendation", "Do not launch VibeFit as another selfie scanner. Launch it as a Korean-inspired personal appearance operating system: one profile that explains what suits the user, turns insight into a small plan and shopping decisions, learns from outcomes, and never rates beauty.")
doc.add_paragraph("The analysis engine is already unusually complete. The next work is the product layer that creates repeat value: goals and constraints, Korean consultation knowledge, product and ingredient compatibility, routines, point-of-decision checks, outcome feedback, expert escalation, calibration evidence, and visible privacy controls.")
doc.add_heading("Best launch wedge", level=2)
doc.add_paragraph("“Your Korean beauty and style consultation, translated into decisions you can use today.” Start with face + personal color + hair + cosmetic skin-wellness guidance. Add body styling next. Keep medical diagnosis, skin-age fear, attractiveness scoring, and golden-ratio ranking outside the product.")

doc.add_heading("1. What VibeFit has already built", level=1)
for item in [
    "Face shape, proportions, symmetry, feature mapping, eyebrow guidance, and a facial overlay.",
    "Undertone, contrast, seasonal color, palettes, makeup color, hair color, and jewelry guidance.",
    "Hair texture/length/density-oriented analysis and haircut recommendations.",
    "Visible cosmetic skin indicators: texture, tone evenness, redness, under-eye appearance, and shine zones.",
    "Body proportions, silhouette guidance, posture hints, necklines, and fit recommendations.",
    "Multi-photo aggregation, scan-quality checks, metric confidence, progress history, share cards, and PDF reports.",
    "A deterministic offline rules engine, with Gemini/Groq used for conversational personalization.",
]: add_bullet(doc, item)
doc.add_paragraph("That foundation matters. Most competitors specialize in a single layer; VibeFit can create one persistent profile across face, color, skin, hair, and body.")

doc.add_heading("2. Why the market is ready", level=1)
doc.add_heading("Appearance analysis is established behavior", level=2)
doc.add_paragraph("Ipsos reports that South Korean Gen Z consumers use personal color, face-shape hair consulting, and skeletal/body-shape consulting for granular identity and style discovery. Bookings can fill immediately even above KRW 100,000. This validates the multi-domain scope and the brand's non-judgmental philosophy.")
doc.add_heading("Global K-beauty creates demand—and overload", level=2)
doc.add_paragraph("Korea's regulator reported about USD 11.4B in cosmetics exports for 2025, up 11.8% year over year, with the United States the largest destination. StyleKorean alone carries 500+ Korean brands across 150 countries. Choice is abundant; confidence is scarce.")
add_callout(doc, "Core unmet job", "Which analysis should I trust, what should I do first, which products fit my constraints, and did they actually work?")

doc.add_heading("3. Competitive signal and whitespace", level=1)
landscape = [
    ("Perfect Corp.", "15-condition skin scans, multi-angle maps, recommendations, tracking", "Raises the bar for repeatability and progress; largely brand/clinic infrastructure"),
    ("Hwahae / Picky", "Ingredients, Korean-user reviews, rankings, matching, community", "Users need explanation and social proof—not a generated shopping list"),
    ("Style DNA", "Color/body/style profile, closet, item checks, daily outfits, shopping", "Analysis is onboarding; daily decisions create retention"),
    ("Myavana", "Hair analysis, goals, products, regimens, tutorials, expert help", "Hair needs a care journey and human escalation"),
    ("Korean bundle apps", "Color, face shape, makeup, hair, body, downloadable reports", "“All-in-one analysis” alone is already commoditizing"),
]
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
hdr = t.rows[0]
set_repeat_table_header(hdr)
for i, txt in enumerate(["Category", "What users can get", "Implication for VibeFit"]):
    hdr.cells[i].text = txt; set_cell_shading(hdr.cells[i], LIGHT)
for a,b,c in landscape:
    cells = t.add_row().cells
    for i, txt in enumerate([a,b,c]): cells[i].text = txt
set_table_geometry(t, [1800, 3300, 4260])

doc.add_heading("4. What people need", level=1)
needs = [
    ("1. Goal-first consultation", "Ask about the outcome, budget, climate, time, sensitivities, maintenance tolerance, preferences, and willingness to change before scanning. Rank the top three actions."),
    ("2. Stable, correctable results", "Show capture quality and which findings are stable or uncertain. Let users correct obvious errors. Publish test-retest performance across devices and diverse users."),
    ("3. Korean expertise, translated", "Build a reviewed taxonomy for personal color, face/bone-structure hair consulting, makeup placement, and body-frame styling. Explain the principle and Korean terminology in plain English."),
    ("4. A minimum viable routine", "Start with a safe baseline, add one change at a time, define frequency, observation window, patch-test guidance, and a stop rule. Avoid ten-product prescriptions."),
    ("5. Independent product compatibility", "Combine profile, ingredients, declared sensitivities, preferences, budget, availability, and current routine. Explain match and mismatch reasons. Separate sponsorship from rank."),
    ("6. Point-of-decision tools", "Check a product, garment, glasses, makeup shade, or hair color before buying; build a Korea shopping list; create bilingual salon and makeup cards."),
    ("7. Outcome learning", "Track comfort, irritation, manageability, confidence, and comparable photos. Connect each outcome to the change introduced; favor weekly/monthly checks over daily score pressure."),
    ("8. Inclusive, non-judgmental guidance", "Use neutral observations and multiple options. Support deeper skin tones, textured hair, varied bodies, men, non-binary users, disabilities, head coverings, and different grooming norms."),
    ("9. Medical boundaries and escalation", "Describe visible cosmetic appearance and wellness habits, not disease or treatment. Route red flags to a dermatologist or qualified hair/scalp professional."),
    ("10. Understandable privacy", "Offer explicit pre-scan consent, local processing where practical, short retention, delete-after-analysis, export/delete, and separate consent for model improvement."),
]
for title, body in needs:
    doc.add_heading(title, level=2)
    doc.add_paragraph(body)

doc.add_heading("5. The product model", level=1)
doc.add_heading("One editable Vibe Profile", level=2)
doc.add_paragraph("Store stable attributes and constraints—personal color, contrast, face geometry, hair characteristics, body proportions, preferences, sensitivities, budget, climate, goals, owned items, and confidence. Every recommendation should state which profile facts drove it.")
doc.add_heading("Four repeatable experiences", level=2)
for text in [
    "Discover: a guided consultation and scan creates the profile and top actions.",
    "Decide: check a product, color, hairstyle, technique, accessory, or garment before spending.",
    "Do: follow a small routine, salon brief, makeup map, or outfit formula.",
    "Learn: track outcomes and refine the profile and recommendations.",
]: add_number(doc, text)
doc.add_heading("Result hierarchy", level=2)
for text in ["Your goal", "Three useful actions now", "Why each is recommended", "Confidence and limitations", "What to avoid or postpone", "Saved plan and next check-in"]: add_number(doc, text)

doc.add_heading("6. Feature priorities", level=1)
doc.add_heading("Build first: market-facing MVP", level=2)
for text in [
    "Goal-and-constraints onboarding.",
    "Unified, editable Vibe Profile with confidence.",
    "Expert-reviewed Korean consultation knowledge base with source/version control.",
    "Action Plan v1: three priorities, minimal skin/hair routine, and salon/style cards.",
    "Product/routine checker v1: ingredients, duplicate actives, declared sensitivities, conflicts, budget, and availability—without medical safety claims.",
    "Outcome loop on every recommendation plus monthly photo comparison.",
    "Trust center: retention, deletion, metric limits, and confidence.",
]: add_number(doc, text)
doc.add_heading("Build next", level=2)
for text in ["Shopping/color/item camera check", "Closet and owned-products inventory", "Weather- and season-aware routines", "Multilingual professional handoff cards", "Credentialed expert review/referrals", "Profile-filtered community reviews with sponsorship labels"]: add_bullet(doc, text)
doc.add_heading("Defer", level=2)
for text in ["Full AR/3D try-on", "Celebrity similarity", "Skin age, golden-ratio, flaw, or attractiveness dashboards", "Medical diagnosis/treatment plans", "A broad affiliate marketplace before trust is established"]: add_bullet(doc, text)

doc.add_heading("7. Positioning and revenue", level=1)
add_callout(doc, "Positioning statement", "For people overwhelmed by beauty and style advice, VibeFit is a Korean-inspired personal appearance companion that turns face, color, skin, hair, and body analysis into explainable routines and confident choices—without rating beauty or pushing one brand.")
doc.add_heading("Initial user", level=2)
doc.add_paragraph("Start with global K-beauty/K-style explorers who already save social content, shop online or plan Korea purchases, and feel overwhelmed by conflicting advice. A sharp test segment is K-beauty beginners with sensitive or combination skin who also want personal color and hair guidance. Validate the segment through interviews.")
doc.add_heading("Revenue ladder", level=2)
for text in [
    "Free: quality-checked scan, basic profile, top three actions, limited checks.",
    "One-time deep consultation: multi-domain report and salon/makeup/shopping cards.",
    "Subscription: ongoing checks, tracking, seasonal updates, owned-item tools, stylist chat.",
    "Affiliate: only after independent ranking; label commission and show alternatives.",
    "Expert layer: paid review or referral with credentials and scope visible.",
]: add_bullet(doc, text)

doc.add_heading("8. 90-day validation plan", level=1)
doc.add_heading("Days 1–30: prove the problem", level=2)
for text in ["Interview 20–30 beginners, enthusiasts, sensitive-skin users, and consultation customers.", "Test three promises: one profile; stop buying wrong products; personal color + hair + routine plan.", "Run expert reviews with a Korean color consultant, hairstylist, makeup artist, formulator/chemist, and dermatologist advisor.", "Audit photo-derived claims for scientific and regulatory defensibility."]: add_bullet(doc, text)
doc.add_heading("Days 31–60: prove comprehension and trust", level=2)
for text in ["Prototype goal intake, profile, top-three action plan, and correction flow.", "Add feedback to every recommendation.", "Build a complete 100–300 product catalog rather than a shallow giant catalog.", "Test repeatability across devices and lighting."]: add_bullet(doc, text)
doc.add_heading("Days 61–90: prove behavior and payment", level=2)
for text in ["Closed beta around “build my beginner Korean routine and personal style profile.”", "Measure saved/followed actions, not scan completion alone.", "Offer a paid one-time deep report before testing subscription."]: add_bullet(doc, text)
add_callout(doc, "North-star metric", "Percentage of activated users who complete a recommended action and report a useful outcome within 30 days.")
doc.add_paragraph("Guardrails: retake disagreement, low-confidence scans, irritation feedback, deletion time, demographic performance gaps, body-image discomfort, and sponsorship influence.")

doc.add_heading("9. Strategic conclusion", level=1)
doc.add_paragraph("The next phase should not be another analyzer. It should connect the existing machinery into a trustworthy loop:")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("UNDERSTAND ME  →  TELL ME WHAT MATTERS  →  HELP ME DECIDE  →  HELP ME DO IT  →  LEARN WHAT WORKED")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(11); r.font.color.rgb = TEAL
doc.add_paragraph("The Korean inspiration is strongest when expressed as granular self-understanding, expert-informed technique, personal color, face/hair/body consulting, gentle routine design, and respectful individuality. VibeFit can win by being broader than a skin scanner, more actionable than a report, more independent than a retailer, and safer and kinder than a beauty-scoring app.")

doc.add_page_break()
doc.add_heading("Sources and limitations", level=1)
sources = [
    ("Ipsos — South Korea 2025: Shifts and Explorations", "https://www.ipsos.com/sites/default/files/ct/publication/documents/2025-01/Ipsos%20Flair_SouthKorea2025_EN.pdf"),
    ("Korean MFDS — 2025 cosmetics trade results", "https://mfds.go.kr/brd/m_1256/view.do?company_cd=&company_nm=&itm_seq_1=0&itm_seq_2=0&multi_itm_seq=0&page=2&seq=23&srchFr=&srchTo=&srchTp=&srchWord="),
    ("Korean MOTIR — K-beauty export channels", "https://english.motir.go.kr/eng/article/EATCLdfa319ada/2394/view"),
    ("Perfect Corp. — AI skin analysis", "https://www.perfectcorp.com/business/products/ai-skin-diagnostic"),
    ("Hwahae Global", "https://www.hwahae.com/en"),
    ("Picky — About", "https://www.gopicky.com/about"),
    ("Style DNA — App Store", "https://apps.apple.com/us/app/style-dna-ai-stylist-closet/id1358319821"),
    ("Myavana", "https://www.myavana.com/"),
    ("American Academy of Dermatology — health and wellness apps", "https://www.aad.org/public/fad/digital-health/apps"),
    ("FDA — mobile medical software functions", "https://www.fda.gov/medical-devices/digital-health-center-excellence/device-software-functions-including-mobile-medical-applications"),
    ("FTC — biometric information policy statement", "https://www.ftc.gov/system/files/ftc_gov/pdf/p225402biometricpolicystatement.pdf"),
]
for label, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    hyperlink(p, label, url)
doc.add_heading("Limitations", level=2)
doc.add_paragraph("This is secondary market research, not a substitute for interviews or willingness-to-pay experiments. Competitor capabilities and validation figures are first-party claims unless independently verified. Community examples were used only as qualitative discovery signals. The launch-country legal analysis must be completed separately.")

# Core properties and final paragraph controls
doc.core_properties.title = "VibeFit Product-Market Research"
doc.core_properties.subject = "Korean-inspired face, skin, hair, color, and body analysis product strategy"
doc.core_properties.author = "VibeFit"
for p in doc.paragraphs:
    p.paragraph_format.widow_control = True
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.widow_control = True

doc.save(OUT)
print(OUT)
